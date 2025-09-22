"""Document service for file processing and management."""

import asyncio
import hashlib
import logging
import mimetypes
import shutil
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, List, Dict, Any, BinaryIO
import uuid

from fastapi import HTTPException, status, UploadFile
from sqlmodel import Session, select
import magic

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
import io

from app.core.config import settings
from app.core.database import get_session
from app.models.document import (
    Document, 
    DocumentCreate, 
    DocumentUpdate, 
    DocumentResponse, 
    DocumentSearchRequest,
    DocumentListResponse,
    ProcessingStatus,
    DocumentType,
    DocumentProcessingStatus
)
from app.models.user import User
from app.services.vector_service import vector_service

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document processing and management."""
    
    def __init__(self, session: Session):
        self.session = session
        self.upload_dir = Path(settings.upload_directory)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def _validate_file(self, file: UploadFile) -> Dict[str, Any]:
        """Validate uploaded file."""
        # Check file size
        if hasattr(file, 'size') and file.size > settings.max_file_size:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File size {file.size} exceeds maximum allowed size {settings.max_file_size}"
            )
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in settings.allowed_file_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type {file_ext} not allowed. Allowed types: {settings.allowed_file_extensions}"
            )
        
        # Detect MIME type
        try:
            # Read first chunk to detect MIME type
            file_header = file.file.read(2048)
            file.file.seek(0)  # Reset file pointer
            
            mime_type = magic.from_buffer(file_header, mime=True)
            
            # Validate MIME type matches extension
            expected_mimes = {
                '.pdf': ['application/pdf'],
                '.docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
                '.txt': ['text/plain', 'text/x-python', 'application/octet-stream']
            }
            
            if file_ext in expected_mimes and mime_type not in expected_mimes[file_ext]:
                # Allow some flexibility for text files
                if file_ext == '.txt' and mime_type.startswith('text/'):
                    pass
                else:
                    logger.warning(f"MIME type mismatch: {mime_type} for extension {file_ext}")
            
        except Exception as e:
            logger.warning(f"Could not detect MIME type: {e}")
            mime_type = mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
        
        return {
            "extension": file_ext,
            "mime_type": mime_type,
            "original_filename": file.filename
        }
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256_hash.update(chunk)
        return sha256_hash.hexdigest()
    
    def _check_duplicate(self, file_hash: str, user_id: int) -> Optional[Document]:
        """Check if document with same hash already exists for user."""
        statement = select(Document).where(
            Document.file_hash == file_hash,
            Document.user_id == user_id,
            Document.processing_status != ProcessingStatus.DELETED
        )
        return self.session.exec(statement).first()
    
    async def upload_document(
        self, 
        file: UploadFile, 
        user: User, 
        document_data: Optional[DocumentCreate] = None
    ) -> DocumentResponse:
        """Upload and process a document."""
        try:
            # Validate file
            file_info = self._validate_file(file)
            
            # Generate unique filename
            file_id = str(uuid.uuid4())
            file_ext = file_info["extension"]
            safe_filename = f"{file_id}{file_ext}"
            file_path = self.upload_dir / safe_filename
            
            # Save file to disk
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            # Calculate file hash
            file_hash = self._calculate_file_hash(file_path)
            
            # Check for duplicates
            existing_doc = self._check_duplicate(file_hash, user.id)
            if existing_doc:
                # Remove uploaded file
                file_path.unlink()
                logger.info(f"Duplicate file detected: {file_hash}")
                return DocumentResponse.from_orm(existing_doc)
            
            # Create document record
            document = Document(
                filename=safe_filename,
                original_filename=file_info["original_filename"],
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                content_type=file_info["mime_type"],
                file_hash=file_hash,
                user_id=user.id,
                processing_status=ProcessingStatus.UPLOADED,
                title=document_data.title if document_data else Path(file.filename).stem,
                description=document_data.description if document_data else None,
                tags=",".join(document_data.tags) if document_data and document_data.tags else None,
                is_confidential=document_data.is_confidential if document_data else True
            )
            
            self.session.add(document)
            self.session.commit()
            self.session.refresh(document)
            
            # Start background processing
            asyncio.create_task(self._process_document_async(document.id))
            
            logger.info(f"Document uploaded: {document.id} by user {user.id}")
            return DocumentResponse.from_orm(document)
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            # Clean up file if it was created
            if 'file_path' in locals() and file_path.exists():
                file_path.unlink()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Document upload failed"
            )
    
    async def _process_document_async(self, document_id: int):
        """Process document asynchronously in the background."""
        try:
            # Get document
            document = self.session.get(Document, document_id)
            if not document:
                logger.error(f"Document {document_id} not found for processing")
                return
            
            # Update status to processing
            document.processing_status = ProcessingStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            self.session.add(document)
            self.session.commit()
            
            # Extract text based on file type
            extracted_text = await self._extract_text(document)
            
            if not extracted_text:
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = "Could not extract text from document"
                self.session.add(document)
                self.session.commit()
                return
            
            # Update document with extracted content
            document.extracted_text = extracted_text
            document.word_count = len(extracted_text.split())
            document.language = self._detect_language(extracted_text)
            
            # Classify document type
            document_type, confidence = await self._classify_document(extracted_text)
            document.document_type = document_type
            document.confidence_score = confidence
            
            # Update status to embedding
            document.processing_status = ProcessingStatus.EMBEDDING
            self.session.add(document)
            self.session.commit()
            
            # Embed document in vector database
            success = await vector_service.embed_document(document, extracted_text)
            
            if success:
                document.processing_status = ProcessingStatus.READY
                document.chroma_collection_id = settings.chroma_collection_name

                # Get chunk count (with graceful handling for when embeddings are disabled)
                try:
                    chunks = await vector_service.get_document_chunks(document.id)
                    document.chunk_count = len(chunks)
                except Exception as chunk_error:
                    logger.warning(f"Could not get chunks for document {document.id}: {chunk_error}")
                    document.chunk_count = 0  # Set to 0 when embeddings are disabled
            else:
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = "Failed to embed document"
            
            document.processing_completed_at = datetime.utcnow()
            self.session.add(document)
            self.session.commit()
            
            logger.info(f"Document processing completed: {document_id} - Status: {document.processing_status}")
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            
            # Update document with error status
            try:
                document = self.session.get(Document, document_id)
                if document:
                    document.processing_status = ProcessingStatus.FAILED
                    document.processing_error = str(e)
                    document.processing_completed_at = datetime.utcnow()
                    self.session.add(document)
                    self.session.commit()
            except Exception as db_error:
                logger.error(f"Error updating document status: {db_error}")
    
    async def _extract_text(self, document: Document) -> Optional[str]:
        """Extract text from document based on file type."""
        try:
            file_path = Path(document.file_path)
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.pdf':
                return await self._extract_text_from_pdf(file_path, document)
            elif file_ext == '.docx':
                return await self._extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                return await self._extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {document.filename}: {e}")
            return None
    
    async def _extract_text_from_pdf(self, file_path: Path, document: Document) -> Optional[str]:
        """Extract text from PDF file with OCR fallback."""
        try:
            text = ""
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                # Try to extract text normally first
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                # Update page count
                document.page_count = page_count
            
            # If no text extracted or very little text, try OCR
            if len(text.strip()) < 100 and settings.enable_ocr:
                logger.info(f"Attempting OCR for document {document.id}")
                ocr_text = await self._extract_text_with_ocr(file_path)
                
                if ocr_text:
                    text = ocr_text
                    document.has_ocr = True
                    document.ocr_confidence = 0.8  # TODO: Get actual confidence from OCR
                    document.processing_status = ProcessingStatus.OCR_PROCESSING
            
            return text.strip() if text else None
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return None
    
    async def _extract_text_from_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            return "\n\n".join(text) if text else None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return None
    
    async def _extract_text_from_txt(self, file_path: Path) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            logger.warning(f"Could not decode text file {file_path} with common encodings")
            return None
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return None
    
    async def _extract_text_with_ocr(self, file_path: Path) -> Optional[str]:
        """Extract text using OCR (Tesseract)."""
        try:
            if not settings.enable_ocr:
                return None
            
            # Configure Tesseract if custom path is provided
            if settings.tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
            
            # For PDF files, convert to images first
            if file_path.suffix.lower() == '.pdf':
                # This would require pdf2image library
                # For now, just return None
                logger.warning("OCR for PDF files requires pdf2image library")
                return None
            else:
                # For image files
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image)
                return text.strip() if text else None
                
        except Exception as e:
            logger.error(f"Error performing OCR on {file_path}: {e}")
            return None
    
    def _detect_language(self, text: str) -> Optional[str]:
        """Detect language of text."""
        # Simple language detection based on common words
        # In production, use a proper language detection library
        
        if not text:
            return None
        
        text_lower = text.lower()
        
        # Check for common English words
        english_words = ['the', 'and', 'or', 'of', 'to', 'in', 'for', 'with', 'on', 'at']
        english_count = sum(1 for word in english_words if word in text_lower)
        
        if english_count >= 3:
            return "en"
        
        return "unknown"
    
    async def _classify_document(self, text: str) -> tuple[DocumentType, float]:
        """Classify document type based on content."""
        # Simple keyword-based classification
        # In production, use a proper ML model
        
        if not text:
            return DocumentType.OTHER, 0.0
        
        text_lower = text.lower()
        
        # Contract keywords
        contract_keywords = ['agreement', 'contract', 'party', 'whereas', 'consideration', 'terms', 'conditions']
        contract_score = sum(1 for keyword in contract_keywords if keyword in text_lower)
        
        # Legal brief keywords
        brief_keywords = ['court', 'plaintiff', 'defendant', 'motion', 'order', 'judgment', 'legal']
        brief_score = sum(1 for keyword in brief_keywords if keyword in text_lower)
        
        # Determine classification
        if contract_score >= 3:
            return DocumentType.CONTRACT, min(0.9, contract_score * 0.15)
        elif brief_score >= 3:
            return DocumentType.LEGAL_BRIEF, min(0.9, brief_score * 0.15)
        else:
            return DocumentType.OTHER, 0.3
    
    async def get_document(self, document_id: int, user: User) -> Optional[DocumentResponse]:
        """Get document by ID."""
        statement = select(Document).where(Document.id == document_id)
        document = self.session.exec(statement).first()
        
        if not document:
            return None
        
        # Check permissions
        if document.user_id != user.id and user.role.value != "admin":
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not enough permissions"
            )
        
        # Update access tracking
        document.last_accessed = datetime.utcnow()
        document.access_count += 1
        self.session.add(document)
        self.session.commit()
        
        return DocumentResponse.from_orm(document)
    
    async def search_documents(
        self, 
        search_request: DocumentSearchRequest, 
        user: User
    ) -> DocumentListResponse:
        """Search documents with filters."""
        try:
            # Build query
            statement = select(Document).where(Document.processing_status != ProcessingStatus.DELETED)
            
            # User filter (non-admins can only see their own documents)
            if user.role.value != "admin":
                statement = statement.where(Document.user_id == user.id)
            elif search_request.user_id:
                statement = statement.where(Document.user_id == search_request.user_id)
            
            # Text search
            if search_request.query:
                statement = statement.where(
                    Document.extracted_text.contains(search_request.query) |
                    Document.title.contains(search_request.query) |
                    Document.description.contains(search_request.query)
                )
            
            # Filters
            if search_request.document_type:
                statement = statement.where(Document.document_type == search_request.document_type)
            
            if search_request.processing_status:
                statement = statement.where(Document.processing_status == search_request.processing_status)
            
            if search_request.date_from:
                statement = statement.where(Document.created_at >= search_request.date_from)
            
            if search_request.date_to:
                statement = statement.where(Document.created_at <= search_request.date_to)
            
            # Sorting
            if search_request.sort_order == "desc":
                statement = statement.order_by(getattr(Document, search_request.sort_by).desc())
            else:
                statement = statement.order_by(getattr(Document, search_request.sort_by))
            
            # Count total results
            total_count = len(self.session.exec(statement).all())
            
            # Pagination
            offset = (search_request.page - 1) * search_request.per_page
            statement = statement.offset(offset).limit(search_request.per_page)
            
            documents = self.session.exec(statement).all()
            
            # Convert to response format
            document_responses = [DocumentResponse.from_orm(doc) for doc in documents]
            
            return DocumentListResponse(
                documents=document_responses,
                total=total_count,
                page=search_request.page,
                per_page=search_request.per_page,
                has_next=offset + search_request.per_page < total_count,
                has_prev=search_request.page > 1
            )
            
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Document search failed"
            )
    
    async def delete_document(self, document_id: int, user: User) -> bool:
        """Delete document (soft delete)."""
        try:
            document = self.session.get(Document, document_id)
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Document not found"
                )
            
            # Check permissions
            if document.user_id != user.id and user.role.value != "admin":
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Not enough permissions"
                )
            
            # Soft delete
            document.processing_status = ProcessingStatus.DELETED
            document.updated_at = datetime.utcnow()
            self.session.add(document)
            self.session.commit()
            
            # Delete from vector database
            await vector_service.delete_document_embeddings(document_id)
            
            logger.info(f"Document deleted: {document_id} by user {user.id}")
            return True
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Document deletion failed"
            )


def get_document_service(session: Session = next(get_session())) -> DocumentService:
    """Get document service instance."""
    return DocumentService(session)